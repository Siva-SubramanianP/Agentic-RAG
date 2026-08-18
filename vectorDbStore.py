import chromadb
from chromadb.utils.embedding_functions import SentenceTransformerEmbeddingFunction
from text_splitter import structure_aware_recursive_chunk
from docx import Document
import fitz
import os
import uuid

CHROMA_PATH = "./demo_CBD"
COLLECTION_NAME = "DEMO"

CHUNK_SIZE = 1000
OVERLAP = 200

def create_chunks(text, chunk_size=CHUNK_SIZE, overlap=OVERLAP):

    if not text.strip():
        return []
    
    chunks = []
    start = 0

    while start < len(text):
        chunk = text[start:start + chunk_size]
        chunks.append(chunk)
        start += chunk_size - overlap

    return chunks


async def vectorDbStoreFunc(documents):

    message = "Document uploaded"
    failureCount = 0
    uploaded_files = []
    all_chunks = []
    all_metadata = []
    all_ids = []

    for file in documents:

        filename = file.filename
        extension = os.path.splitext(filename)[1].lower()

        content = await file.read()

        pages = []
        text = ""

        try:
            if extension == ".pdf":

                document = fitz.open(stream=content, filetype="pdf")

                for page_number, page in enumerate(document):
                    text = page.get_text().strip()

                    if text:
                        pages.append({
                            "page": page_number + 1,
                            "text": text
                        })

                document.close()


            elif extension == ".txt":

                text = content.decode("utf-8")

                pages.append({
                    "page": 1,
                    "text": text
                })


            elif extension == ".docx":

                temp_filename = f"temp_{filename}"

                with open(temp_filename, "wb") as f:
                    f.write(content)

                doc = Document(temp_filename)

                for paragraph in doc.paragraphs:
                    if paragraph.text.strip():
                        text += paragraph.text + "\n"


                for table_index, table in enumerate(doc.tables):

                    text += f"\n--- Table {table_index+1} ---\n"

                    for row in table.rows:
                        row_data = []

                        for cell in row.cells:
                            row_data.append(cell.text.strip())

                        text += " | ".join(row_data) + "\n"


                pages.append({
                    "page": 1,
                    "text": text
                })

                os.remove(temp_filename)


            else:
                failureCount += 1
                uploaded_files.append({
                    "filename" : filename,
                    "pages" : 0
                })
                continue


            uploaded_files.append({
                "filename": filename,
                "pages": pages
            })

            for page in pages:

                # page_chunks = create_chunks(page["text"])
                chunks = structure_aware_recursive_chunk(
                    markdown_text=page["text"],
                    document_name=filename,
                    chunk_size=500,
                    chunk_overlap=75,
                )


                for index, chunk in enumerate(chunks):

                    chroma_id = str(uuid.uuid4())

                    all_chunks.append(chunk["text"])
                    all_ids.append(chroma_id)

                    metadata = {
                        "filename": filename
                    }

                    all_metadata.append(metadata)


        except Exception as e:
            failureCount += 1
            print(f"Failed processing {filename}: {e}")

    
    #print("Stored", len(chunks), "chunks")

    client = chromadb.PersistentClient(
        path=CHROMA_PATH
    )

    # ef = embedding_functions.DefaultEmbeddingFunction()

    

    ef = SentenceTransformerEmbeddingFunction(model_name="all-MiniLM-L6-v2")

    collection = client.get_or_create_collection(
        name=COLLECTION_NAME,
        embedding_function=ef
    )

    if not all_chunks:
        return "Upload failed", uploaded_files, failureCount

    existing = collection.get(
        include=["metadatas"]
    )

    existing_filenames = set()

    for metadata in existing["metadatas"]:
        if metadata and "filename" in metadata:
            existing_filenames.add(metadata["filename"])

    filtered_ids = []
    filtered_chunks = []
    filtered_metadata = []

    for chunk_id, chunk, metadata in zip(
        all_ids,
        all_chunks,
        all_metadata):

        if metadata["filename"] not in existing_filenames:
            filtered_ids.append(chunk_id)
            filtered_chunks.append(chunk)
            filtered_metadata.append(metadata)

    if filtered_chunks:
        collection.add(
            ids=filtered_ids,
            documents=filtered_chunks,
            metadatas=filtered_metadata
        )
    else:
        message = "No new documents to add"
        failureCount += 1

    return message, uploaded_files , failureCount

    #print("Data stored successfully.")